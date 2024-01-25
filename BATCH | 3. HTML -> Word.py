from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tkinter import Tk, filedialog
import os
import base64
from io import BytesIO
from PIL import Image

def convert_base64_to_rgb(base64_data):
    # Decode base64 data and convert to RGB
    image_data = base64.b64decode(base64_data)
    image = Image.open(BytesIO(image_data))
    image_rgb = image.convert("RGB")
    return image_rgb

def save_image_as_jpeg(image_rgb, output_path):
    # Save the image as JPEG
    image_rgb.save(output_path, format="JPEG")

def handle_list_item(paragraph, element):
    # Handle list items with bullet points
    run = paragraph.add_run('• ')
    font = run.font
    font.size = Pt(18)  # Font size 18
    paragraph.add_run(element.get_text())

def handle_table(doc, element):
    # Handle tables
    table = doc.add_table(rows=len(element.find_all('tr')), cols=max(len(row.find_all(['td', 'th'])) for row in element.find_all('tr')))
    for row_index, row in enumerate(element.find_all('tr')):
        for col_index, cell in enumerate(row.find_all(['td', 'th'])):
            table.cell(row_index, col_index).text = cell.get_text()

def handle_figure(doc, element, image_counter):
    # Handle figure elements
    if element.find('img') and element.find('img').get('src').startswith('data:image'):
        base64_data = element.find('img')['src'].split(',', 1)[1]
        image_rgb = convert_base64_to_rgb(base64_data)
        temp_image_path = f"temp_image_{image_counter}.jpeg"
        save_image_as_jpeg(image_rgb, temp_image_path)
        doc.add_picture(temp_image_path, width=Inches(5))
        os.remove(temp_image_path)

def handle_math(doc, element):
    # Handle math elements
    run = doc.add_paragraph().add_run(element.get_text())
    font = run.font
    font.size = Pt(18)  # Font size 18

def html_to_word(input_path, output_path):
    # Load HTML content
    with open(input_path, 'r', encoding='utf-8') as html_file:
        html_content = html_file.read()

    # Create a Word document
    doc = Document()

    # Set default font for the entire document
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(18)

    # Parse HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Mapping HTML headings to Word styles
    heading_styles = {
        'h1': 'Heading1',
        'h2': 'Heading2',
        'h3': 'Heading3',
        'h4': 'Heading4',
        'h5': 'Heading5',
        'h6': 'Heading6',
    }

    # Counter for image placeholders
    image_counter = 1

    # Initialize last_line
    last_line = ""

    # Iterate through HTML tags and convert to Word format
    for element in soup.descendants:
        if element.name == 'p':
            paragraph = doc.add_paragraph()

            # Set paragraph alignment if specified in HTML
            if element.get('style'):
                styles = element['style'].split(';')
                for style in styles:
                    if 'text-align' in style:
                        alignment = style.split(':')[-1].strip()
                        if alignment == 'left':
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        elif alignment == 'center':
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        elif alignment == 'right':
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Add text to the paragraph
            run = paragraph.add_run(element.get_text())
            font = run.font
            font.size = Pt(18)  # Font size 18

            # Update last_line
            last_line = element.get_text()

        elif element.name in heading_styles:
            # Handle HTML headings
            heading_level = int(element.name[1])
            heading_style = heading_styles[element.name]
            heading = doc.add_paragraph(element.get_text(), style=heading_style)
            heading.style.font.bold = True  # Make heading text bold
            heading.style.font.size = Pt(2 * (18 + 2 * (1 - heading_level)))  # Double and adjust font size based on heading level

            # Update last_line
            last_line = element.get_text()

        elif element.name in ['strong', 'b']:
            # Handle bold text
            bold_text = element.get_text()

            # Check if the last line is the same and delete it
            if last_line.strip() == bold_text.strip():
                # Delete the last paragraph
                doc.paragraphs[-1].clear()

            run = doc.add_paragraph().add_run(bold_text)
            font = run.font
            font.bold = True
            font.size = Pt(18)  # Font size 18

            # Update last_line
            last_line = bold_text

        elif element.name == 'img' and element.get('src') and element.get('src').startswith('data:image'):
            # Handle image elements with base64 data
            base64_data = element['src'].split(',', 1)[1]  # Extract base64 data
            image_rgb = convert_base64_to_rgb(base64_data)

            # Save the image as JPEG
            temp_image_path = f"temp_image_{image_counter}.jpeg"
            save_image_as_jpeg(image_rgb, temp_image_path)

            # Replace the image placeholder with the actual image
            doc.add_picture(temp_image_path, width=Inches(5))

            os.remove(temp_image_path)  # Remove the temporary image file

            # Increment the image counter
            image_counter += 1

        elif element.name == 'li':
            # Handle list items
            handle_list_item(doc.add_paragraph(), element)

        elif element.name == 'table':
            # Handle tables
            handle_table(doc, element)

        elif element.name == 'figure':
            # Handle figures
            handle_figure(doc, element, image_counter)

        elif element.name == 'math':
            # Handle math elements
            handle_math(doc, element)

    # Save the Word document
    doc.save(output_path)

    print(f"Conversion successful. Word document saved at: {output_path}")

def select_folder():
    root = Tk()
    root.withdraw()  # Hide the main window

    folder_path = filedialog.askdirectory(title="Select a folder of HTML files", initialdir=os.getcwd())
    return folder_path

def batch_process_folder(input_folder):
    output_folder = os.path.join(input_folder, "HTMLtoWORD")
    os.makedirs(output_folder, exist_ok=True)

    for file_name in os.listdir(input_folder):
        if file_name.endswith(".html") or file_name.endswith(".htm"):
            try:
                input_path = os.path.join(input_folder, file_name)
                output_path = os.path.join(output_folder, file_name.replace(".html", ".docx").replace(".htm", ".docx"))
                html_to_word(input_path, output_path)
            except Exception as e:
                print(f"Error processing {file_name}: {e}")

if __name__ == "__main__":
    input_folder = select_folder()

    if input_folder:
        batch_process_folder(input_folder)

import os
from docx import Document
from tkinter import Tk, filedialog

def get_largest_font_size(paragraph):
    return max((run.font.size.pt for run in paragraph.runs if run.font.size), default=12)

def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        if not doc.paragraphs:
            return None
        largest_font_paragraph = max(doc.paragraphs, key=get_largest_font_size)
        return largest_font_paragraph.text
    except Exception as e:
        print(f"Error processing {docx_path}: {e}")
        return None

def process_files(folder_path):
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            docx_path = os.path.join(folder_path, filename)
            extracted_text = extract_text_from_docx(docx_path)
            
            if extracted_text is not None:
                # Take the first sentence as the new filename
                new_filename = extracted_text.split('.')[0]
                
                # Ensure the new filename contains legal symbols for macOS
                new_filename = "".join(c if c.isalnum() or c in {' ', '_', '-'} else '_' for c in new_filename)
                
                # Ensure the new filename length doesn't exceed macOS limits
                new_filename = new_filename[:255]  # Limit to 255 characters
                
                # Rename the file
                new_path = os.path.join(folder_path, new_filename + ".docx")
                os.rename(docx_path, new_path)

if __name__ == "__main__":
    # Use Tkinter to select a folder
    root = Tk()
    root.withdraw()
    folder_path = filedialog.askdirectory(title="Select Folder with Word Documents")

    if folder_path:
        process_files(folder_path)
        print("Batch processing completed.")
    else:
        print("No folder selected. Exiting.")
